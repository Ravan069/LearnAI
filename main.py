import streamlit as st
from PyPDF2 import PdfReader
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
import google.generativeai as genai
from langchain_google_genai import ChatGoogleGenerativeAI
from docx import Document

# API Key Fetching
gemini_api_key = st.secrets.get("GOOGLE_API_KEY")
if gemini_api_key is None:
    raise ValueError("GOOGLE_API_KEY not found in environment variables")
genai.configure(api_key=gemini_api_key)

# PDF to Text
def get_pdf_text(pdf_docs):
    text = ""
    for pdf in pdf_docs:
        pdf_reader = PdfReader(pdf)
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text

# Text to Chunk
def get_text_chunks(raw_text):
    text_splitter = CharacterTextSplitter(
        separator="\n",
        chunk_size=1000,
        chunk_overlap=200
    )
    return text_splitter.split_text(raw_text)

# Chunks to Embeddings
def get_vector(chunks):
    if not chunks:
        return None
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", api_key=gemini_api_key)
    return FAISS.from_texts(texts=chunks, embedding=embeddings)

# Convo Chain
def conversation_chain():
    template = """
    You are the Teacher, help the students by answering the question asked, which are related to the document uploaded. The following are guidelines:
    - Answer the question in 2500 words.
    - Draw diagrams using simple geometry.
    - Use the data from the document to answer the question.
    - Provide a small summary explaining the answer; the goal is to make students understand the 2500-word answer.
    Context: \n{context}\n
    Question: \n{question}\n
    Answer:
    """
    model_instance = ChatGoogleGenerativeAI(model="gemini-1.5-flash", api_key=gemini_api_key)
    prompt = PromptTemplate(template=template, input_variables=["context", "question"])
    return load_qa_chain(model_instance, chain_type="stuff", prompt=prompt), model_instance

# User Question
def user_question(question, db, chain, raw_text):
    if db is None:
        return "Please upload and process a PDF first."

    docs = db.similarity_search(question, k=5)
    response = chain.invoke(
        {"input_documents": docs, "question": question, "context": raw_text},
        return_only_outputs=True
    )
    return response.get("output_text")

# Save Conversation to Word Document
def save_conversation_to_docx(messages):
    doc = Document()
    doc.add_heading("Conversation", level=1)

    for message in messages:
        role = "User" if message["role"] == "user" else "Assistant"
        doc.add_heading(f"{role}:", level=2)
        doc.add_paragraph(message["content"])

    byte_stream = BytesIO()
    doc.save(byte_stream)
    byte_stream.seek(0)
    return byte_stream

# Main
def main():
    # Set the page configuration
    st.set_page_config(page_title="Learn AI 😊", page_icon=":books:", layout="wide")
    st.header("Learn AI 😊")

    # Initialize session state variables specific to this page
    if "messages_chatbot_2" not in st.session_state:
        st.session_state.messages_chatbot_2 = []
    if "vector_store_chatbot_2" not in st.session_state:
        st.session_state.vector_store_chatbot_2 = None
    if "chain_chatbot_2" not in st.session_state:
        st.session_state.chain_chatbot_2 = None
    if "raw_text_chatbot_2" not in st.session_state:
        st.session_state.raw_text_chatbot_2 = None
        
    with st.sidebar:
        st.subheader("Upload Your PDFs")
        pdf_docs = st.file_uploader("Choose PDF files", accept_multiple_files=True, type="pdf")

        if st.button("Process PDF"):
            if not pdf_docs:
                st.error("Please upload at least one PDF file.")
            else:
                with st.spinner("Processing..."):
                    raw_text = get_pdf_text(pdf_docs)
                    chunks = get_text_chunks(raw_text)
                    vector_store = get_vector(chunks)
                    chain, _ = conversation_chain()

                    if vector_store and chain and raw_text:
                        st.session_state.vector_store_chatbot_2 = vector_store
                        st.session_state.chain_chatbot_2 = chain
                        st.session_state.raw_text_chatbot_2 = raw_text
                        st.success("PDF processed successfully.")  

        # Provide a download button for the conversation
        

    for message in st.session_state.messages_chatbot_2:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # Input for user questions
    if st.session_state.vector_store_chatbot_2 and st.session_state.chain_chatbot_2 and st.session_state.raw_text_chatbot_2:
        user_query = st.chat_input("Ask your question:")
        if user_query:
            with st.chat_message("user"):
                st.markdown(user_query)
            st.session_state.messages_chatbot_2.append({"role": "user", "content": user_query})

            response = user_question(user_query, st.session_state.vector_store_chatbot_2, st.session_state.chain_chatbot_2, st.session_state.raw_text_chatbot_2)
            if response:
                st.session_state.messages_chatbot_2.append({"role": "assistant", "content": response})
                with st.chat_message("assistant"):
                    st.markdown(response)
                    
        if st.session_state.messages_chatbot_2:
                docx_file = save_conversation_to_docx(st.session_state.messages_chatbot_2)
                st.download_button(
                label="Download Conversation as .docx",
                data=docx_file,
                file_name="conversation.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )               
                         

if __name__ == '__main__':
    main()
